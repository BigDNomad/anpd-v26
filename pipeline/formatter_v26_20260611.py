#!/usr/bin/env python3
"""
v22 Formatter — Assembles chapters into a single .docx manuscript.
KDP 5x8 paperback standard.

Enhancements over 20260303:
    - B7: First-paragraph indent suppression (after chapter title and scene breaks)
    - B8: Title page, copyright, static TOC, alternating headers, footer page numbers
    - Even/odd headers via XML (python-docx doesn't natively support)
    - Title page suppression (no header/footer on page 1)

Usage:
    python3 formatter_20260305.py \
        --chapters-dir /root/anpd-system/v22/series/billy_gamble/books/book_04/output/chapters \
        --output /root/anpd-system/v22/series/billy_gamble/books/book_04/output/Switchback_v22.docx \
        --title "Switchback" --author "David Lee Corley"

Copyright (c) 2025 Endeavor Publishing LLC
Created: 2026-03-05
"""

import os
import sys
import re
import json
import argparse
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


CHAPTER_NUMBERS = {
    1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five",
    6: "Six", 7: "Seven", 8: "Eight", 9: "Nine", 10: "Ten",
    11: "Eleven", 12: "Twelve", 13: "Thirteen", 14: "Fourteen", 15: "Fifteen",
    16: "Sixteen", 17: "Seventeen", 18: "Eighteen", 19: "Nineteen", 20: "Twenty",
    21: "Twenty-One", 22: "Twenty-Two", 23: "Twenty-Three",
    24: "Twenty-Four", 25: "Twenty-Five", 26: "Twenty-Six",
    27: "Twenty-Seven", 28: "Twenty-Eight", 29: "Twenty-Nine", 30: "Thirty",
    31: "Thirty-One", 32: "Thirty-Two", 33: "Thirty-Three",
    34: "Thirty-Four", 35: "Thirty-Five", 36: "Thirty-Six",
    37: "Thirty-Seven", 38: "Thirty-Eight", 39: "Thirty-Nine", 40: "Forty",
    41: "Forty-One", 42: "Forty-Two", 43: "Forty-Three",
    44: "Forty-Four", 45: "Forty-Five", 46: "Forty-Six",
    47: "Forty-Seven", 48: "Forty-Eight", 49: "Forty-Nine", 50: "Fifty",
    51: "Fifty-One", 52: "Fifty-Two", 53: "Fifty-Three",
    54: "Fifty-Four", 55: "Fifty-Five", 56: "Fifty-Six",
    57: "Fifty-Seven", 58: "Fifty-Eight", 59: "Fifty-Nine", 60: "Sixty",
    61: "Sixty-One", 62: "Sixty-Two", 63: "Sixty-Three",
    64: "Sixty-Four", 65: "Sixty-Five", 66: "Sixty-Six",
    67: "Sixty-Seven", 68: "Sixty-Eight", 69: "Sixty-Nine", 70: "Seventy",
    71: "Seventy-One"
}


def _add_run(paragraph, text, font_name='Garamond', font_size=Pt(11),
             bold=False, italic=False):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    return run


def _setup_headers_footers(section, title, author):
    """Set up alternating headers and footer page numbers via XML.

    Even pages: BOOK TITLE (italic, centered, Garamond 9pt)
    Odd pages: AUTHOR NAME (centered, Garamond 9pt)
    Footer: Centered page number on all pages (Garamond 9pt)
    Title page: No header/footer (first page different)
    """
    sectPr = section._sectPr

    # Enable even/odd headers and title page suppression
    sectPr.append(parse_xml('<w:titlePg ' + nsdecls('w') + '/>'))

    # --- First page header (empty — title page suppression) ---
    first_hdr = parse_xml(
        '<w:hdr ' + nsdecls('w', 'r') + '>'
        '  <w:p><w:pPr><w:jc w:val="center"/></w:pPr></w:p>'
        '</w:hdr>'
    )

    # --- Even page header: BOOK TITLE (italic) ---
    even_hdr = parse_xml(
        '<w:hdr ' + nsdecls('w', 'r') + '>'
        '  <w:p>'
        '    <w:pPr><w:jc w:val="center"/></w:pPr>'
        '    <w:r>'
        '      <w:rPr>'
        '        <w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/>'
        '        <w:sz w:val="18"/>'
        '        <w:i/>'
        '      </w:rPr>'
        '      <w:t>' + _xml_escape(title.upper()) + '</w:t>'
        '    </w:r>'
        '  </w:p>'
        '</w:hdr>'
    )

    # --- Odd (default) page header: AUTHOR NAME ---
    odd_hdr = parse_xml(
        '<w:hdr ' + nsdecls('w', 'r') + '>'
        '  <w:p>'
        '    <w:pPr><w:jc w:val="center"/></w:pPr>'
        '    <w:r>'
        '      <w:rPr>'
        '        <w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/>'
        '        <w:sz w:val="18"/>'
        '      </w:rPr>'
        '      <w:t>' + _xml_escape(author) + '</w:t>'
        '    </w:r>'
        '  </w:p>'
        '</w:hdr>'
    )

    # --- Footer: centered page number (all pages except first) ---
    ftr = parse_xml(
        '<w:ftr ' + nsdecls('w', 'r') + '>'
        '  <w:p>'
        '    <w:pPr><w:jc w:val="center"/></w:pPr>'
        '    <w:r>'
        '      <w:rPr>'
        '        <w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/>'
        '        <w:sz w:val="18"/>'
        '      </w:rPr>'
        '      <w:fldChar w:fldCharType="begin"/>'
        '    </w:r>'
        '    <w:r>'
        '      <w:rPr>'
        '        <w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/>'
        '        <w:sz w:val="18"/>'
        '      </w:rPr>'
        '      <w:instrText xml:space="preserve"> PAGE </w:instrText>'
        '    </w:r>'
        '    <w:r>'
        '      <w:rPr>'
        '        <w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/>'
        '        <w:sz w:val="18"/>'
        '      </w:rPr>'
        '      <w:fldChar w:fldCharType="end"/>'
        '    </w:r>'
        '  </w:p>'
        '</w:ftr>'
    )

    # First page footer (empty — no page number on title page)
    first_ftr = parse_xml(
        '<w:ftr ' + nsdecls('w', 'r') + '>'
        '  <w:p><w:pPr><w:jc w:val="center"/></w:pPr></w:p>'
        '</w:ftr>'
    )

    # Add relationships and reference in section properties
    # We need to add these to the document's part
    part = section.part

    # Add header/footer parts
    first_hdr_rId = part.relate_to(part._new_hdr_ftr_part('header', first_hdr), 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header')
    even_hdr_rId = part.relate_to(part._new_hdr_ftr_part('header', even_hdr), 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header')
    odd_hdr_rId = part.relate_to(part._new_hdr_ftr_part('header', odd_hdr), 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header')
    ftr_rId = part.relate_to(part._new_hdr_ftr_part('footer', ftr), 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer')
    first_ftr_rId = part.relate_to(part._new_hdr_ftr_part('footer', first_ftr), 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer')

    # Add references in sectPr
    sectPr.append(parse_xml(f'<w:headerReference {nsdecls("w", "r")} w:type="first" r:id="{first_hdr_rId}"/>'))
    sectPr.append(parse_xml(f'<w:headerReference {nsdecls("w", "r")} w:type="even" r:id="{even_hdr_rId}"/>'))
    sectPr.append(parse_xml(f'<w:headerReference {nsdecls("w", "r")} w:type="default" r:id="{odd_hdr_rId}"/>'))
    sectPr.append(parse_xml(f'<w:footerReference {nsdecls("w", "r")} w:type="default" r:id="{ftr_rId}"/>'))
    sectPr.append(parse_xml(f'<w:footerReference {nsdecls("w", "r")} w:type="first" r:id="{first_ftr_rId}"/>'))


def _setup_headers_footers_simple(doc, title, author):
    """Simplified header/footer setup using python-docx API + minimal XML.

    This approach avoids the _new_hdr_ftr_part internal API.
    """
    section = doc.sections[0]
    sectPr = section._sectPr

    # Enable different first page (title page suppression)
    section.different_first_page_header_footer = True

    # Enable even/odd headers in document settings
    settings = doc.settings.element
    even_odd = settings.find(qn('w:evenAndOddHeaders'))
    if even_odd is None:
        settings.append(parse_xml('<w:evenAndOddHeaders ' + nsdecls('w') + '/>'))

    # --- Default (odd) header: author name ---
    odd_header = section.header
    odd_header.is_linked_to_previous = False
    odd_para = odd_header.paragraphs[0]
    odd_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = odd_para.add_run(author)
    run.font.name = 'Garamond'
    run.font.size = Pt(9)

    # --- Even header: book title (ALL CAPS, not italic) ---
    even_header = section.even_page_header
    even_header.is_linked_to_previous = False
    even_para = even_header.paragraphs[0]
    even_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = even_para.add_run(title.upper())
    run.font.name = 'Garamond'
    run.font.size = Pt(9)
    run.italic = False

    # --- First page header: empty ---
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    # Already empty

    # --- Default footer: page number ---
    footer = section.footer
    footer.is_linked_to_previous = False
    ftr_para = footer.paragraphs[0]
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add PAGE field code via XML
    fld_begin = parse_xml(
        '<w:r ' + nsdecls('w') + '>'
        '  <w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="18"/></w:rPr>'
        '  <w:fldChar w:fldCharType="begin"/>'
        '</w:r>'
    )
    fld_code = parse_xml(
        '<w:r ' + nsdecls('w') + '>'
        '  <w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="18"/></w:rPr>'
        '  <w:instrText xml:space="preserve"> PAGE </w:instrText>'
        '</w:r>'
    )
    fld_end = parse_xml(
        '<w:r ' + nsdecls('w') + '>'
        '  <w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="18"/></w:rPr>'
        '  <w:fldChar w:fldCharType="end"/>'
        '</w:r>'
    )
    ftr_para._p.append(fld_begin)
    ftr_para._p.append(fld_code)
    ftr_para._p.append(fld_end)

    # --- Even footer: same page number ---
    even_footer = section.even_page_footer
    even_footer.is_linked_to_previous = False
    eftr_para = even_footer.paragraphs[0]
    eftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    efld_begin = parse_xml(
        '<w:r ' + nsdecls('w') + '>'
        '  <w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="18"/></w:rPr>'
        '  <w:fldChar w:fldCharType="begin"/>'
        '</w:r>'
    )
    efld_code = parse_xml(
        '<w:r ' + nsdecls('w') + '>'
        '  <w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="18"/></w:rPr>'
        '  <w:instrText xml:space="preserve"> PAGE </w:instrText>'
        '</w:r>'
    )
    efld_end = parse_xml(
        '<w:r ' + nsdecls('w') + '>'
        '  <w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="18"/></w:rPr>'
        '  <w:fldChar w:fldCharType="end"/>'
        '</w:r>'
    )
    eftr_para._p.append(efld_begin)
    eftr_para._p.append(efld_code)
    eftr_para._p.append(efld_end)

    # --- First page footer: empty ---
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    # Already empty


def _xml_escape(text):
    """Escape text for XML embedding."""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def _add_paragraph_with_italic(doc, para_text, first_line_indent=True):
    """Add a paragraph with italic markdown handling.

    Args:
        doc: Document instance
        para_text: Text with *italic* markers
        first_line_indent: Whether to apply first-line indent (B7: suppress for first para)

    Returns:
        The added paragraph
    """
    p = doc.add_paragraph()
    parts = re.split(r'(\*[^*]+\*)', para_text)
    for part in parts:
        if part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(11)
        run.font.name = 'Garamond'

    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    return p


def format_manuscript(chapters_dir, output_path, title="The Eighteen",
                       author="David Lee Corley",
                       copyright_text="Copyright © 2026 Endeavor Publishing LLC"):
    """Assemble chapter .md files into a formatted .docx manuscript.

    Features:
        - Title page with title, author, copyright
        - Static table of contents
        - Alternating headers (even: title, odd: author)
        - Footer page numbers
        - First-paragraph indent suppression (B7)
        - 5x8 KDP paperback format
    """

    doc = Document()

    # --- Page setup: 5x8 KDP paperback ---
    # KDP-safe mirror margins. With mirrorMargins ON (injected into
    # settings.xml below), Word reinterprets left_margin as the INSIDE
    # (binding/gutter) edge and right_margin as the OUTSIDE edge, and
    # alternates them across the spine. Inside 0.75" clears KDP's gutter
    # minimum for every 5x8 page-count tier (0.625" req at 300+pp).
    # Without mirrorMargins, a flat left=0.75/right=0.5 leaves only 0.5"
    # on the binding edge of even pages -> KDP "insufficient gutter"
    # rejection (the Mandate/btd001 bounce, 2026-05-24).
    for section in doc.sections:
        section.page_width = Inches(5)
        section.page_height = Inches(8)
        section.top_margin = Inches(0.60)
        section.bottom_margin = Inches(0.60)
        section.left_margin = Inches(0.75)   # INSIDE edge under mirrorMargins
        section.right_margin = Inches(0.60)  # OUTSIDE edge under mirrorMargins

    # --- Default font ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.15

    # --- Headers and footers (B8) ---
    _setup_headers_footers_simple(doc, title, author)

    # --- Title page ---
    for _ in range(8):
        doc.add_paragraph('')

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title.upper())
    run.font.size = Pt(24)
    run.font.name = 'Garamond'
    run.bold = True

    doc.add_paragraph('')

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run(author)
    run.font.size = Pt(14)
    run.font.name = 'Garamond'

    # --- Copyright page ---
    doc.add_page_break()
    doc.add_paragraph('')
    doc.add_paragraph('')

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(copyright_text)
    run.font.size = Pt(9)
    run.font.name = 'Garamond'

    cp2 = doc.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp2.add_run('All rights reserved.')
    run.font.size = Pt(9)
    run.font.name = 'Garamond'

    # --- Collect and sort chapter files ---
    chapter_files = []
    for fname in sorted(os.listdir(chapters_dir)):
        if not fname.endswith('.md'):
            continue
        # Match ch01_one.md pattern (skip attempt and raw files)
        m = re.match(r'ch(\d+)_.+\.md$', fname)
        if m and '_attempt' not in fname and '_raw' not in fname and '_original' not in fname:
            ch_num = int(m.group(1))
            chapter_files.append((ch_num, os.path.join(chapters_dir, fname)))

    chapter_files.sort(key=lambda x: x[0])

    # --- Static Table of Contents (B8) ---
    doc.add_page_break()

    for _ in range(3):
        doc.add_paragraph('')

    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run('Contents')
    run.font.size = Pt(16)
    run.font.name = 'Garamond'
    run.bold = True

    doc.add_paragraph('')

    for ch_num, _ in chapter_files:
        ch_word = CHAPTER_NUMBERS.get(ch_num, str(ch_num))
        toc_entry = doc.add_paragraph()
        toc_entry.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_entry.add_run(f"Chapter {ch_word}")
        run.font.size = Pt(11)
        run.font.name = 'Garamond'
        toc_entry.paragraph_format.space_after = Pt(4)

    # --- Chapters ---
    total_words = 0
    chapter_stats = []

    for ch_num, ch_path in chapter_files:
        with open(ch_path) as f:
            text = f.read().strip()

        word_count = len(text.split())
        total_words += word_count
        chapter_stats.append({
            'chapter': ch_num,
            'title': f"Chapter {CHAPTER_NUMBERS.get(ch_num, str(ch_num))}",
            'words': word_count,
            'file': os.path.basename(ch_path),
        })

        # --- New page for each chapter ---
        doc.add_page_break()

        # --- Chapter heading ---
        # Skip any markdown heading in the text
        lines = text.split('\n')
        start_idx = 0

        for i, line in enumerate(lines):
            stripped = line.strip()
            # Skip markdown headings like "# Chapter One"
            if stripped.startswith('#'):
                start_idx = i + 1
                continue
            # Skip "Chapter One", "Chapter 26", etc. as plain text heading
            if re.match(r'^Chapter\s+\S.*$', stripped, re.IGNORECASE):
                start_idx = i + 1
                continue
            break

        # Add chapter title with formatting — ALL CAPS, not italic
        ch_word = CHAPTER_NUMBERS.get(ch_num, str(ch_num))
        expected_title = f"Chapter {ch_word}".upper()

        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_para.paragraph_format.space_before = Pt(72)
        run = heading_para.add_run(expected_title)
        run.font.size = Pt(16)
        run.font.name = 'Garamond'
        run.bold = True
        run.italic = False

        # Spacing after title
        doc.add_paragraph('')
        doc.add_paragraph('')

        # --- Chapter body ---
        body_text = '\n'.join(lines[start_idx:]).strip()

        # B7: Track whether next paragraph should suppress indent
        suppress_indent = True  # First paragraph after chapter title

        for para_text in body_text.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            # Scene break
            if para_text == '* * *':
                break_para = doc.add_paragraph()
                break_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = break_para.add_run('* * *')
                run.font.size = Pt(11)
                run.font.name = 'Garamond'
                break_para.paragraph_format.space_before = Pt(12)
                break_para.paragraph_format.space_after = Pt(12)
                suppress_indent = True  # Next paragraph after scene break
                continue

            # Regular paragraph — B7: suppress indent for first para after title/break
            _add_paragraph_with_italic(
                doc, para_text,
                first_line_indent=(not suppress_indent)
            )
            suppress_indent = False

    # --- KDP mirror margins (settings.xml) ---
    # python-docx has no API for mirrorMargins; inject the element directly.
    # Must be present for the left/right margins above to alternate across
    # the spine. Idempotent: only adds if not already present.
    settings_el = doc.settings.element
    if settings_el.find(qn('w:mirrorMargins')) is None:
        settings_el.append(parse_xml('<w:mirrorMargins ' + nsdecls('w') + '/>'))

    # --- Save ---
    doc.save(output_path)

    return {
        'output_path': output_path,
        'total_words': total_words,
        'chapters': len(chapter_stats),
        'chapter_stats': chapter_stats,
    }


def main():
    parser = argparse.ArgumentParser(description='v23 Formatter (20260406)')
    parser.add_argument('--chapters-dir', required=True)
    # Standalone mode: caller gives an explicit output file path.
    parser.add_argument('--output', default=None,
                        help='Explicit output .docx path (standalone mode)')
    parser.add_argument('--title', default=None)
    parser.add_argument('--author', default=None,
                        help='Author name (overrides intake pen_name if set)')
    parser.add_argument('--intake', default=None,
                        help='Path to intake JSON; reads pen_name and title')
    # Orchestrator mode: caller gives a dir + book metadata; formatter owns
    # filename construction per the handle_format contract:
    #   {NN}_{Title}_{YYYYMMDD_HHMM}.docx
    parser.add_argument('--output-dir', default=None,
                        help='Output directory (orchestrator mode); filename is built')
    parser.add_argument('--book-number', default=None,
                        help='Book number for filename (orchestrator mode)')
    parser.add_argument('--book-title', default=None,
                        help='Book title (orchestrator mode; alias for --title)')
    parser.add_argument('--author-name', default=None,
                        help='Author name (orchestrator mode; alias for --author)')
    args = parser.parse_args()

    # Resolve author + title from any of the accepted aliases / intake.
    author = args.author or args.author_name
    title = args.title or args.book_title
    if (author is None or title is None) and args.intake:
        with open(args.intake) as f:
            intake = json.load(f)
        if author is None:
            author = intake.get('pen_name')
        if title is None:
            title = intake.get('title')
    if author is None:
        parser.error('No author: supply --author/--author-name or --intake with pen_name')
    if title is None:
        parser.error('No title: supply --title/--book-title or --intake with title')

    # Resolve output path: explicit --output wins; else build from --output-dir.
    output_path = args.output
    if output_path is None:
        if args.output_dir is None:
            parser.error('No output: supply --output (file) or --output-dir (directory)')
        import datetime as _dt
        # {NN}_{Title}_{YYYYMMDD_HHMM}.docx — NN zero-padded if numeric.
        try:
            nn = f"{int(args.book_number):02d}" if args.book_number is not None else "00"
        except (TypeError, ValueError):
            nn = str(args.book_number)
        safe_title = re.sub(r'[^\w\- ]', '', title).strip().replace(' ', '_')
        stamp = _dt.datetime.now().strftime('%Y%m%d_%H%M')
        fname = f"{nn}_{safe_title}_{stamp}.docx"
        os.makedirs(args.output_dir, exist_ok=True)
        output_path = os.path.join(args.output_dir, fname)

    result = format_manuscript(
        chapters_dir=args.chapters_dir,
        output_path=output_path,
        title=title,
        author=author,
    )

    print(f"\nManuscript formatted: {result['output_path']}")
    print(f"Total words: {result['total_words']:,}")
    print(f"Chapters: {result['chapters']}")
    print(f"\nPer-chapter word counts:")
    for ch in result['chapter_stats']:
        print(f"  {ch['title']}: {ch['words']:,} words")


if __name__ == '__main__':
    main()
