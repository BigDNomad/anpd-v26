#!/usr/bin/env python3
# RULE 1: Do not modify this ANPD pipeline file without explicit permission from Dave.
# RULE 2: If this component fails, write STOP_REPORT.json and halt immediately.
# RULE 3: Do not skip, comment out, or work around this component.
# RULE 4: Do not modify master_controller.py to bypass this component.
# RULE 5: Report the exact error. Do not attempt a fix.
"""
ANPD V24 — Manuscript Summarizer (v3)

Narrative-style review of a finished V24 manuscript. Companion to
synopsis_summarizer.py.

Reads either a final .docx (post-formatter) or a markdown source.
Produces a scene-cluster-resolution review (~3 pages) walking the
manuscript at higher granularity than the synopsis review — instead of
one paragraph per chapter, scene-cluster paragraphs covering 1-2
related scenes each.

Same action thriller rubric as synopsis_summarizer: strict action
definition (gunfire, hand-to-hand, explosions, chases, escapes), Light/
Medium/Heavy strength tiers, climax check, gap-and-proposal pairs.

Differs from the synopsis review in two ways: (1) higher resolution —
scene-cluster paragraphs instead of chapter paragraphs, (2) reviews
*delivered* execution rather than planned structure. The manuscript
review answers "did the manuscript deliver what the synopsis promised,
and where did action actually land in the prose."

Per White Paper §3.6: analytical-mode prompt — judgments against the
rubric. Surfaces concerns for Dave to act on; not pushed back into the
running pipeline.

Output: {book_dir}/out/reviews/{Title}_man_rev_{YYYYMMDD_HHMM}.docx
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

from llm_client import call_llm

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write(
        "ERROR: 'python-docx' package not installed. Run: pip install python-docx\n"
    )
    sys.exit(1)


DEFAULT_MODEL = "claude-sonnet-4-5"
# Manuscript reviews produce more output than synopsis reviews because the
# review walks the manuscript at scene-cluster resolution — 60-80 cluster
# paragraphs covering 1-2 scenes each across the original 100-scene structure,
# plus opening / action_distribution / climax_check / delivery_gaps /
# structural_concerns sections. Sized to 28000 to give 2× headroom over the
# expected output budget.
MAX_OUTPUT_TOKENS = 28000


# ──────────────────────────────────────────────────────────────────────────────
# Helpers (consistent with synopsis_summarizer)
# ──────────────────────────────────────────────────────────────────────────────


def strip_leading_article(title: str) -> str:
    return re.sub(r"^(the|a|an)\s+", "", title, flags=re.IGNORECASE)


def slugify_title(title: str) -> str:
    title = strip_leading_article(title)
    title = re.sub(r"[\'\"\:,]", "", title)
    title = re.sub(r"[^\w\s-]", "", title)
    title = re.sub(r"[\s-]+", "_", title.strip())
    return title


def resolve_title(
    book_config_path: Optional[Path],
    intake_path: Optional[Path],
    fallback: str = "Untitled",
) -> str:
    for path in (book_config_path, intake_path):
        if path and path.exists():
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
                title = data.get("title", "").strip()
                if title:
                    return title
            except (json.JSONDecodeError, OSError):
                continue
    return fallback


def resolve_model(
    series_config_path: Optional[Path],
    cli_override: Optional[str],
) -> str:
    if cli_override:
        return cli_override
    if series_config_path and series_config_path.exists():
        try:
            cfg = json.loads(series_config_path.read_text(encoding="utf-8"))
            return (
                cfg.get("model_manuscript_summarizer")
                or cfg.get("model_manuscript_audit_whole")
                or DEFAULT_MODEL
            )
        except (json.JSONDecodeError, OSError) as e:
            sys.stderr.write(
                f"WARNING: Could not read series_config: {e}\n"
                f"Falling back to default model {DEFAULT_MODEL}.\n"
            )
    return DEFAULT_MODEL


def resolve_output_dir(manuscript_path: Path, override: Optional[Path]) -> Path:
    if override:
        return override
    parent = manuscript_path.resolve().parent
    # Walk up to find a V24 book directory (parent of out/ or work/).
    for candidate in [parent] + list(parent.parents)[:5]:
        if (candidate / "out").is_dir():
            return candidate / "out" / "reviews"
        if candidate.name == "out":
            return candidate.parent / "out" / "reviews"
        if candidate.name == "work":
            return candidate.parent / "out" / "reviews"
    # Fallback for V23-or-earlier manuscripts with non-V24 layouts.
    return parent / "reviews"


def extract_manuscript_text(path: Path) -> str:
    """Read manuscript text from .docx, .md, or .txt.

    Fails loudly per §2.1 on unsupported format or read error.
    """
    suffix = path.suffix.lower()
    if suffix in (".md", ".txt"):
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        try:
            doc = Document(str(path))
        except Exception as e:
            sys.stderr.write(f"ERROR: Could not open .docx {path}: {e}\n")
            sys.exit(1)
        return "\n".join(p.text for p in doc.paragraphs)
    sys.stderr.write(
        f"ERROR: Unsupported manuscript format: {suffix}. "
        f"Supported: .docx, .md, .txt\n"
    )
    sys.exit(1)


# ──────────────────────────────────────────────────────────────────────────────
# Action thriller rubric (same as synopsis_summarizer for consistency)
# ──────────────────────────────────────────────────────────────────────────────

ACTION_THRILLER_RUBRIC = """# Action Thriller Rubric

You are reviewing a finished thriller manuscript as a story editor
specialized in action thrillers. Your job is to help the author
identify where the manuscript delivered the story well, where action
landed effectively, and where the prose underdelivered on the
synopsis's structural promises.

## Action Scene — Strict Definition

A scene qualifies as ACTION if and only if the prose contains one or
more of:

- Hand-to-hand fighting
- Gunfire (firing or being fired upon)
- Edged weapons or other lethal-force weapons in active use
- Explosions or controlled destruction
- Chase scenes (protagonist pursuing or being pursued, on foot or vehicle)
- Escape scenes (breaking out, evading capture, slipping a closing net)

The unifying test: the scene raises the reader's adrenaline as written.

NOT action, regardless of plot importance:
- Driving from one location to another, even if plot-relevant
- Surveillance, reconnaissance, observation
- Planning, briefing, intelligence-gathering
- Tense conversation or interrogation, even with high stakes
- Emotional confrontation without physical danger

## Strength Tiers

LIGHT. A single physical beat. Brief. One person engaged on the
protagonist's side. Resolved quickly, often without damage taken.

MEDIUM. Multiple physical beats in sequence, or sustained engagement
with a single opponent. Stakes are personal but contained.

HEAVY. Sustained, multi-axis action. Multiple combatants, escalating
intensity, real damage taken or dealt, environmental hazard or chase
pressure layered on top of the immediate fight. Reads like a setpiece.

## Climax

The climax is normally chapter 24 of 25. It should be the HEAVIEST
action scene in the book. A medium or light climax is a flag.

The climax should also contain:
- Direct confrontation between protagonist and antagonist
- Real cost to the protagonist (damage, loss, sacrifice)
- Environmental pressure beyond the immediate fight
- Resolution driven by the protagonist, not external rescue

## Distribution Expectations

- No more than two consecutive chapters should pass without action
- At least one heavy action scene should land before the first twist (around scene 25)
- Front slow + back-loaded action is a genre failure
- All five action types should appear at least once across the book
- The same action type should not appear in two consecutive action scenes

## Manuscript-Specific Failure Modes

In addition to synopsis-level concerns, watch for delivery problems:

- Action scenes that the synopsis promised but the prose didn't deliver
  (a "gunfight" reduced to two paragraphs of summary)
- Action scenes that started strong and trailed off into description
- Chapters where pacing on the page sags despite plot momentum
- Established weapons or skills that are never used in action
- Repetition of the same kind of action beat across chapters
- Climax delivered through dialogue or revelation rather than action
"""


# ──────────────────────────────────────────────────────────────────────────────
# Prompt
# ──────────────────────────────────────────────────────────────────────────────


SUMMARIZER_PROMPT_TEMPLATE = """{rubric}

# YOUR TASK

Read the finished manuscript below and produce a narrative-style review
in JSON format. The review walks the manuscript at scene-cluster
resolution — paragraphs covering 1-2 related scenes each, grouped
naturally by chapter. Identify where action lands well, where it
underdelivers, where pacing sags, and where the manuscript's execution
diverges from what an action thriller needs.

This is a review of *delivered* execution, not planned structure. The
synopsis review (run earlier) caught structural problems. Your job is
to catch delivery problems in the actual prose.

Tone: a working story editor talking to the author about the finished
draft. Honest, specific, constructive. No hedging, no padding. If the
climax was promised heavy and delivered medium, say so. If a scene
started strong and trailed off, name it.

# MANUSCRIPT

{manuscript_text}

# OUTPUT FORMAT

Return a JSON object with these keys:

- "opening": 4-5 sentences orienting the reader. Who is the protagonist,
  what is the situation, what is the spine of the story as the
  manuscript actually delivers it. Mention chapter count, approximate
  word count, and any structural deviation from a 25-chapter / 100-scene
  thriller you observe.

- "scene_clusters": array of objects walking the manuscript at HIGHER
  RESOLUTION than a chapter-by-chapter walk. Cover the full manuscript
  using 60-80 cluster entries total — each cluster covers 1-2 scenes
  from the original 100-scene structure. A 25-chapter book typically
  produces 2-4 clusters per chapter (most chapters split into 2-3
  clusters; action-dense or twist-bearing chapters may split to 4).
  Each object has:
    - "chapter": integer chapter number this cluster falls in
    - "label": short identifier for the cluster, scoped to the cluster's
      content ("Ch 5 — alley test", "Ch 12 — compound infiltration",
      "Ch 24 — bodyguard gunfight", "Ch 24 — fleeing escape"). Don't
      number scenes if the manuscript doesn't number them; just describe
      the cluster.
    - "narrative": 2-4 sentences describing what happens in this
      cluster AND evaluating it. When action lands, name the strength
      tier inline ("a heavy gunfight", "a light takedown that the prose
      undersells"). When you flag a delivery concern, do so directly
      ("the chase reads as summary, not as scene", "this confrontation
      was promised heavy in the synopsis, the prose delivers medium").
      Use "TWIST 1", "TWIST 2", "TWIST 3", and "CLIMAX" markers where
      relevant. If a scene cluster is pure setup with no action, you
      can flag it briefly without dwelling.

  Cluster distribution discipline: do NOT collapse multiple chapters
  into a single cluster. Do NOT skip chapters. Each chapter must contain
  at least 2 clusters covering its scenes; chapters with 4+ scenes
  should produce 2-3 clusters; chapters with twists, climax beats, or
  multiple action scenes should produce 3-4 clusters.

- "action_distribution": one paragraph (5-7 sentences) summarizing how
  action lands across the manuscript. Include: (a) approximate count of
  action scenes by your strict definition vs. total scene count, (b)
  which chapters carry the heaviest action, (c) which stretches sag,
  (d) whether action is balanced or front/back-loaded, (e) which of the
  five action types appeared and which were absent.

- "climax_check": 3-4 sentences evaluating the climax against the
  rubric. State the climax's strength tier explicitly. If the climax is
  medium or light, propose specifically how the prose could be revised
  to elevate it to heavy. If the climax was structurally heavy in the
  synopsis but the prose underdelivered, name where the gap opened.

- "delivery_gaps": array of objects identifying where the prose
  underdelivered on the synopsis's structural promises. Each object:
    - "where": chapter and scene-cluster identifier
    - "concern": one sentence describing what underdelivered
    - "proposal": one sentence proposing a specific revision
  Include 3-7 of these. Focus on highest-leverage delivery problems.

- "structural_concerns": array of objects for problems that affect the
  manuscript at book-level rather than scene-level. Same shape:
    - "concern": one sentence
    - "proposal": one sentence
  Include 1-4 of these. Examples: established weapons never used,
  protagonist takes no damage, action variety too narrow, climax
  resolved through dialogue.

Return ONLY valid JSON. No preamble, no markdown fences, no trailing
commentary."""


def build_prompt(manuscript_text: str) -> str:
    return SUMMARIZER_PROMPT_TEMPLATE.format(
        rubric=ACTION_THRILLER_RUBRIC,
        manuscript_text=manuscript_text,
    )


# ──────────────────────────────────────────────────────────────────────────────
# Model invocation
# ──────────────────────────────────────────────────────────────────────────────


def call_model(prompt: str, model: str) -> str:
    """Call Claude API with streaming via llm_client.

    Streaming keeps the connection alive for long generations (>1 min)
    that can time out with non-streaming requests.
    """
    try:
        response = call_llm(
            provider="anthropic",
            model=model,
            system="You are a story editor reviewing a manuscript.",
            user=prompt,
            max_tokens=MAX_OUTPUT_TOKENS,
            stream=True,
        )
    except Exception as e:
        sys.stderr.write(
            f"ERROR: Streaming call to model {model} failed: {e}\n"
        )
        sys.exit(1)

    if not response.text:
        sys.stderr.write(
            f"ERROR: Empty response from model {model}\n"
        )
        sys.exit(1)
    return response.text


def parse_response(raw: str) -> dict:
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```\s*$", "", cleaned)

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as e:
        sys.stderr.write(
            f"ERROR: Model returned non-JSON content. Parse error: {e}\n"
            f"Raw response (first 500 chars):\n{raw[:500]}\n"
        )
        sys.exit(1)

    required_keys = {
        "opening",
        "scene_clusters",
        "action_distribution",
        "climax_check",
        "delivery_gaps",
        "structural_concerns",
    }
    missing = required_keys - set(data.keys())
    if missing:
        sys.stderr.write(
            f"ERROR: Model response missing required keys: {sorted(missing)}\n"
            f"Got keys: {sorted(data.keys())}\n"
        )
        sys.exit(1)

    if not isinstance(data["scene_clusters"], list) or not data["scene_clusters"]:
        sys.stderr.write("ERROR: 'scene_clusters' must be a non-empty array.\n")
        sys.exit(1)
    for i, cl in enumerate(data["scene_clusters"]):
        required = ("chapter", "label", "narrative")
        if not isinstance(cl, dict) or not all(k in cl for k in required):
            sys.stderr.write(
                f"ERROR: scene_clusters[{i}] missing required fields. "
                f"Expected {{'chapter', 'label', 'narrative'}}; got {cl}\n"
            )
            sys.exit(1)

    for key in ("delivery_gaps", "structural_concerns"):
        if not isinstance(data[key], list):
            sys.stderr.write(f"ERROR: '{key}' must be an array.\n")
            sys.exit(1)
        for i, item in enumerate(data[key]):
            if not isinstance(item, dict):
                sys.stderr.write(f"ERROR: {key}[{i}] is not an object.\n")
                sys.exit(1)
            if "concern" not in item or "proposal" not in item:
                sys.stderr.write(
                    f"ERROR: {key}[{i}] missing 'concern' or 'proposal'. "
                    f"Got: {item}\n"
                )
                sys.exit(1)

    return data


# ──────────────────────────────────────────────────────────────────────────────
# Document rendering
# ──────────────────────────────────────────────────────────────────────────────


def render_docx(data: dict, title: str, output_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title block.
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Manuscript Review")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(
        f"Generated {datetime.now().strftime('%B %d, %Y at %H:%M')}"
    )
    date_run.font.size = Pt(9)

    doc.add_paragraph()

    # Opening paragraph.
    opening_para = doc.add_paragraph(data["opening"])
    opening_para.paragraph_format.space_after = Pt(8)

    # Scene-cluster narratives, grouped by chapter.
    current_chapter = None
    for cl in data["scene_clusters"]:
        # Insert a chapter divider when we move to a new chapter.
        if cl["chapter"] != current_chapter:
            current_chapter = cl["chapter"]
            ch_heading = doc.add_paragraph()
            ch_run = ch_heading.add_run(f"Chapter {current_chapter}")
            ch_run.bold = True
            ch_run.font.size = Pt(11)
            ch_heading.paragraph_format.space_before = Pt(6)
            ch_heading.paragraph_format.space_after = Pt(2)

        para = doc.add_paragraph()
        label_run = para.add_run(f"{cl['label']}. ")
        label_run.bold = True
        para.add_run(cl["narrative"])
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.left_indent = Inches(0.15)

    doc.add_paragraph()

    # Action distribution.
    dist_heading = doc.add_paragraph()
    dist_run = dist_heading.add_run("Action distribution")
    dist_run.bold = True
    dist_run.font.size = Pt(11)
    dist_para = doc.add_paragraph(data["action_distribution"])
    dist_para.paragraph_format.space_after = Pt(6)

    # Climax check.
    climax_heading = doc.add_paragraph()
    climax_run = climax_heading.add_run("Climax check")
    climax_run.bold = True
    climax_run.font.size = Pt(11)
    climax_para = doc.add_paragraph(data["climax_check"])
    climax_para.paragraph_format.space_after = Pt(6)

    # Delivery gaps.
    if data["delivery_gaps"]:
        gaps_heading = doc.add_paragraph()
        gaps_run = gaps_heading.add_run("Delivery gaps")
        gaps_run.bold = True
        gaps_run.font.size = Pt(11)

        for gp in data["delivery_gaps"]:
            where = gp.get("where", "")
            concern_para = doc.add_paragraph()
            concern_run = concern_para.add_run(
                f"Concern{f' ({where})' if where else ''}. "
            )
            concern_run.bold = True
            concern_para.add_run(gp["concern"])
            concern_para.paragraph_format.space_after = Pt(2)

            prop_para = doc.add_paragraph()
            prop_run = prop_para.add_run("Proposal. ")
            prop_run.bold = True
            prop_para.add_run(gp["proposal"])
            prop_para.paragraph_format.space_after = Pt(6)

    # Structural concerns.
    if data["structural_concerns"]:
        struct_heading = doc.add_paragraph()
        struct_run = struct_heading.add_run("Structural concerns")
        struct_run.bold = True
        struct_run.font.size = Pt(11)

        for sc in data["structural_concerns"]:
            concern_para = doc.add_paragraph()
            concern_run = concern_para.add_run("Concern. ")
            concern_run.bold = True
            concern_para.add_run(sc["concern"])
            concern_para.paragraph_format.space_after = Pt(2)

            prop_para = doc.add_paragraph()
            prop_run = prop_para.add_run("Proposal. ")
            prop_run.bold = True
            prop_para.add_run(sc["proposal"])
            prop_para.paragraph_format.space_after = Pt(6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────


def main() -> int:
    parser = argparse.ArgumentParser(
        description="ANPD V24 manuscript summarizer (v3) — narrative-style "
        "scene-cluster review of finished manuscripts."
    )
    parser.add_argument("--manuscript", type=Path, required=True)
    parser.add_argument("--series-config", type=Path, default=None)
    parser.add_argument("--book-config", type=Path, default=None)
    parser.add_argument("--intake", type=Path, default=None)
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--model", type=str, default=None)
    args = parser.parse_args()

    if not args.manuscript.exists():
        sys.stderr.write(f"ERROR: Manuscript not found: {args.manuscript}\n")
        return 1

    manuscript_text = extract_manuscript_text(args.manuscript).strip()
    if not manuscript_text:
        sys.stderr.write(f"ERROR: Manuscript is empty: {args.manuscript}\n")
        return 1

    title = resolve_title(args.book_config, args.intake)
    model = resolve_model(args.series_config, args.model)
    output_dir = resolve_output_dir(args.manuscript, args.output_dir)

    slug = slugify_title(title)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = output_dir / f"{slug}_man_rev_{timestamp}.docx"

    word_count = len(manuscript_text.split())
    sys.stdout.write(f"Manuscript:  {args.manuscript}\n")
    sys.stdout.write(f"Word count:  {word_count:,}\n")
    sys.stdout.write(f"Title:       {title}\n")
    sys.stdout.write(f"Model:       {model}\n")
    sys.stdout.write(f"Output:      {output_path}\n")
    sys.stdout.write("Calling model...\n")

    prompt = build_prompt(manuscript_text)
    raw_response = call_model(prompt, model)
    data = parse_response(raw_response)
    render_docx(data, title, output_path)

    sys.stdout.write(f"DONE. Review written to: {output_path}\n")
    return 0


if __name__ == "__main__":
    sys.exit(main())
