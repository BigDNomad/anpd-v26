#!/usr/bin/env python3
# RULE 1: Do not modify this ANPD pipeline file without explicit permission from Dave.
# RULE 2: If this component fails, write STOP_REPORT.json and halt immediately.
# RULE 3: Do not skip, comment out, or work around this component.
# RULE 4: Do not modify master_controller.py to bypass this component.
# RULE 5: Report the exact error. Do not attempt a fix.
"""
ANPD V25 — Synopsis Summarizer (v4)

Produces two reports from a synopsis:
1. Quality Review (2-3 pages) — structural assessment for go/no-go decision
2. Story Summary (2-3 pages) — plot/arc/turns summary for quick reference

Two LLM calls, two docx outputs per run.

Output: {book_dir}/out/reviews/{Title}_quality_review_{YYYYMMDD_HHMM}.docx
        {book_dir}/out/reviews/{Title}_story_summary_{YYYYMMDD_HHMM}.docx
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

from llm_client import call_llm

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write(
        "ERROR: 'python-docx' package not installed. Run: pip install python-docx\n"
    )
    sys.exit(1)


DEFAULT_MODEL = "claude-sonnet-4-5"
# Each report targets 800-1200 words (~3000-4200 tokens). 8000 gives comfortable headroom.
MAX_OUTPUT_TOKENS = 8000


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────


def strip_leading_article(title: str) -> str:
    return re.sub(r"^(the|a|an)\s+", "", title, flags=re.IGNORECASE)


def slugify_title(title: str) -> str:
    title = strip_leading_article(title)
    title = re.sub(r"[\'\"\:,]", "", title)
    title = re.sub(r"[^\w\s-]", "", title)
    title = re.sub(r"[\s-]+", "_", title.strip())
    return title


def resolve_title(
    book_config_path: Optional[Path],
    intake_path: Optional[Path],
    fallback: str = "Untitled",
) -> str:
    for path in (book_config_path, intake_path):
        if path and path.exists():
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
                title = data.get("title", "").strip()
                if title:
                    return title
            except (json.JSONDecodeError, OSError):
                continue
    return fallback


def resolve_model(
    series_config_path: Optional[Path],
    cli_override: Optional[str],
) -> str:
    if cli_override:
        return cli_override
    if series_config_path and series_config_path.exists():
        try:
            cfg = json.loads(series_config_path.read_text(encoding="utf-8"))
            return (
                cfg.get("model_synopsis_summarizer")
                or cfg.get("model_synopsis_generation")
                or DEFAULT_MODEL
            )
        except (json.JSONDecodeError, OSError) as e:
            sys.stderr.write(
                f"WARNING: Could not read series_config: {e}\n"
                f"Falling back to default model {DEFAULT_MODEL}.\n"
            )
    return DEFAULT_MODEL


def resolve_output_dir(synopsis_path: Path, override: Optional[Path]) -> Path:
    if override:
        return override
    parent = synopsis_path.resolve().parent
    if parent.name == "work":
        book_dir = parent.parent
        return book_dir / "out" / "reviews"
    return parent / "reviews"


# ──────────────────────────────────────────────────────────────────────────────
# Prompt builders — two reports
# ──────────────────────────────────────────────────────────────────────────────


def build_quality_review_prompt(synopsis_text: str) -> str:
    """Prompt for a 2-3 page compressed structural quality review."""
    return f"""You are a story editor specialized in action thrillers. Read the
synopsis below and produce a compressed QUALITY REVIEW for the author.

PURPOSE: Help the author decide whether to PROCEED to manuscript generation,
PROCEED WITH CAUTION, or DO NOT PROCEED.

# SYNOPSIS

{synopsis_text}

# OUTPUT FORMAT

Return a JSON object with these keys:

- "verdict": one of "PROCEED", "PROCEED WITH CAUTION", or "DO NOT PROCEED"
- "verdict_rationale": 1-2 sentences explaining the verdict
- "pacing": 2-4 sentences on overall pacing (front-loaded? back-loaded? balanced?)
- "twists": 2-4 sentences evaluating twist positions (expected at ~25%, ~50%, ~75% of scenes). Name the twist scenes and whether they land at the right structural positions.
- "action_distribution": 2-4 sentences on action scene count, spacing, variety. Flag stretches with no action or clustering.
- "final_battle": 2-3 sentences on the climax. Is it the heaviest action scene? Direct protagonist/antagonist confrontation? Real cost?
- "resolution": 1-2 sentences on emotional resolution after the climax.
- "gaps": array of 2-5 objects, each with "concern" (1 sentence) and "proposal" (1 sentence). Highest-leverage structural changes only.

CONSTRAINTS:
- Total response MUST be 800-1200 words. Conciseness is mandatory.
- Bullet-dense, no narrative walkthrough, no scene-by-scene commentary.
- Analytical and blunt. No hedging, no padding.
- Return ONLY valid JSON. No preamble, no markdown fences."""


def build_story_summary_prompt(synopsis_text: str) -> str:
    """Prompt for a 2-3 page plot/arc/turns summary."""
    return f"""You are a story editor. Read the synopsis below and produce a
STORY SUMMARY — a compressed narrative of the plot for quick reference.

PURPOSE: Let someone recall the full story arc without re-reading the
synopsis. Like an expanded book jacket that covers the entire plot.

# SYNOPSIS

{synopsis_text}

# OUTPUT FORMAT

Return a JSON object with these keys:

- "setup": 2-4 sentences. Who is the protagonist? What is the world state at the opening? What is the inciting event?
- "rising_action": 3-5 sentences. First act complications. Key alliances formed. Antagonist forces introduced.
- "midpoint": 2-3 sentences. The midpoint turn that changes the protagonist's understanding or situation.
- "crisis": 3-5 sentences. Escalating problems after midpoint. Losses. Betrayals. The lowest point.
- "climax": 2-4 sentences. The final confrontation. What happens, who is involved, what it costs.
- "resolution": 2-3 sentences. What has changed. Where the characters end up. Any open threads.
- "protagonist_arc": 1-2 sentences. How the protagonist changes from beginning to end.
- "central_conflict": 1 sentence. The core dramatic question of the book.

CONSTRAINTS:
- Total response MUST be 800-1200 words.
- Confident, vivid, compressed prose. Not analysis — storytelling.
- Name characters. Name locations. Be specific.
- Do NOT include quality assessments, gap callouts, structural analysis, or scene numbers.
- Return ONLY valid JSON. No preamble, no markdown fences."""


def build_beat_summary_prompt(synopsis_text: str) -> str:
    """Prompt for a one-line-per-scene beat summary."""
    return f"""You are a story editor. Read the synopsis below and produce a
BEAT SUMMARY — exactly one entry per scene, for quick structural reference.

# SYNOPSIS

{synopsis_text}

# OUTPUT FORMAT

Return a JSON object with a single key:

- "scenes": array of objects, one per scene in the synopsis. Each object has:
    - "number": integer scene number
    - "title": scene title from the synopsis header
    - "type": scene TYPE tag (ACTION, NON-ACTION, SUSPENSE, MIXED)
    - "pov": POV character name
    - "beat": one-clause description of what happens (15-25 words, action verb, no adjectives unless load-bearing)

CONSTRAINTS:
- Produce EXACTLY one entry per scene in the synopsis. Do not merge scenes. Do not skip scenes.
- If the synopsis has 100 scenes, your output must have exactly 100 entries.
- Description must be a single clause — no commentary, no judgments, no quality language.
- No headers, no preamble, no closing — just the JSON object.
- Return ONLY valid JSON. No markdown fences."""


# ──────────────────────────────────────────────────────────────────────────────
# Model invocation
# ──────────────────────────────────────────────────────────────────────────────


def call_model(prompt: str, model: str, max_tokens: int = MAX_OUTPUT_TOKENS) -> str:
    """Call Claude API with streaming via llm_client."""
    try:
        response = call_llm(
            provider="anthropic",
            model=model,
            system="You are a story editor reviewing a synopsis.",
            user=prompt,
            max_tokens=max_tokens,
            stream=True,
        )
    except Exception as e:
        sys.stderr.write(f"ERROR: Call to model {model} failed: {e}\n")
        sys.exit(1)

    if not response.text:
        sys.stderr.write(f"ERROR: Empty response from model {model}\n")
        sys.exit(1)
    return response.text


def parse_json_response(raw: str, label: str = "response") -> dict:
    """Parse JSON from LLM response. Exits 1 on failure."""
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```\s*$", "", cleaned)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        sys.stderr.write(
            f"FATAL: {label} — non-JSON content. Parse error: {e}\n"
            f"Response length: {len(raw)} chars\n"
            f"Raw response (first 500 chars):\n{raw[:500]}\n"
        )
        dump_name = f"summarizer_failure_{label}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        dump_path = Path(dump_name)
        try:
            dump_path.write_text(raw, encoding="utf-8")
            sys.stderr.write(f"Raw response saved to: {dump_path}\n")
        except OSError:
            pass
        sys.exit(1)


# ──────────────────────────────────────────────────────────────────────────────
# Document rendering
# ──────────────────────────────────────────────────────────────────────────────


def _render_docx_base(title: str, subtitle: str) -> Document:
    """Create a Document with standard page setup and title block."""
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(
        f"Generated {datetime.now().strftime('%B %d, %Y at %H:%M')}"
    )
    date_run.font.size = Pt(9)
    doc.add_paragraph()
    return doc


def render_quality_review_docx(data: dict, title: str, output_path: Path) -> None:
    """Render the quality review JSON into a docx."""
    doc = _render_docx_base(title, "Quality Review")

    # Verdict block
    verdict_para = doc.add_paragraph()
    v_run = verdict_para.add_run(f"VERDICT: {data.get('verdict', 'UNKNOWN')}")
    v_run.bold = True
    v_run.font.size = Pt(13)
    if data.get("verdict_rationale"):
        doc.add_paragraph(data["verdict_rationale"]).paragraph_format.space_after = Pt(8)

    # Sections
    sections = [
        ("Pacing", "pacing"),
        ("Twists", "twists"),
        ("Action Distribution", "action_distribution"),
        ("Final Battle", "final_battle"),
        ("Resolution", "resolution"),
    ]
    for heading, key in sections:
        if data.get(key):
            h = doc.add_paragraph()
            h_run = h.add_run(heading)
            h_run.bold = True
            p = doc.add_paragraph(data[key])
            p.paragraph_format.space_after = Pt(6)

    # Gaps
    gaps = data.get("gaps", [])
    if gaps:
        h = doc.add_paragraph()
        h_run = h.add_run("Gaps and Proposals")
        h_run.bold = True
        for gp in gaps:
            cp = doc.add_paragraph()
            cr = cp.add_run("Concern: ")
            cr.bold = True
            cp.add_run(gp.get("concern", ""))
            cp.paragraph_format.space_after = Pt(2)
            pp = doc.add_paragraph()
            pr = pp.add_run("Proposal: ")
            pr.bold = True
            pp.add_run(gp.get("proposal", ""))
            pp.paragraph_format.space_after = Pt(6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def render_story_summary_docx(data: dict, title: str, output_path: Path) -> None:
    """Render the story summary JSON into a docx."""
    doc = _render_docx_base(title, "Story Summary")

    sections = [
        ("Setup", "setup"),
        ("Rising Action", "rising_action"),
        ("Midpoint", "midpoint"),
        ("Crisis", "crisis"),
        ("Climax", "climax"),
        ("Resolution", "resolution"),
        ("Protagonist Arc", "protagonist_arc"),
        ("Central Conflict", "central_conflict"),
    ]
    for heading, key in sections:
        if data.get(key):
            h = doc.add_paragraph()
            h_run = h.add_run(heading)
            h_run.bold = True
            p = doc.add_paragraph(data[key])
            p.paragraph_format.space_after = Pt(6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def render_beat_summary_docx(data: dict, title: str, output_path: Path) -> None:
    """Render the beat summary JSON into a docx — one entry per scene."""
    doc = _render_docx_base(title, "Beat Summary")

    scenes = data.get("scenes", [])
    for sc in scenes:
        para = doc.add_paragraph()
        header_run = para.add_run(
            f"Scene {sc.get('number', '?')} \u2014 {sc.get('title', '')} "
            f"[{sc.get('type', '?')} / {sc.get('pov', '?')}]"
        )
        header_run.bold = True
        header_run.font.size = Pt(10)
        beat_para = doc.add_paragraph(f"  {sc.get('beat', '')}")
        beat_para.paragraph_format.space_after = Pt(4)
        beat_para.style.font.size = Pt(10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ──────────────────────────────────────────────────────────────────────────────
# Library callable
# ──────────────────────────────────────────────────────────────────────────────


def summarize_synopsis(
    synopsis_path: Path,
    intake_path: Optional[Path] = None,
    output_dir: Optional[Path] = None,
    model: Optional[str] = None,
    max_tokens: int = MAX_OUTPUT_TOKENS,
) -> dict:
    """Run both reports and return paths. For use by other pipeline components.

    Returns {"status": "success"|"failure", "quality_review_path": Path,
             "story_summary_path": Path, "error": str|None}
    """
    synopsis_text = synopsis_path.read_text(encoding="utf-8").strip()
    title = resolve_title(None, intake_path)
    effective_model = model or DEFAULT_MODEL
    effective_output_dir = output_dir or resolve_output_dir(synopsis_path, None)

    slug = slugify_title(title)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    quality_path = effective_output_dir / f"{slug}_quality_review_{timestamp}.docx"
    summary_path = effective_output_dir / f"{slug}_story_summary_{timestamp}.docx"
    beat_path = effective_output_dir / f"{slug}_beat_summary_{timestamp}.docx"

    try:
        # Quality Review
        qr_prompt = build_quality_review_prompt(synopsis_text)
        qr_raw = call_model(qr_prompt, effective_model, max_tokens=max_tokens)
        qr_data = parse_json_response(qr_raw, "quality_review")
        render_quality_review_docx(qr_data, title, quality_path)

        # Story Summary
        ss_prompt = build_story_summary_prompt(synopsis_text)
        ss_raw = call_model(ss_prompt, effective_model, max_tokens=max_tokens)
        ss_data = parse_json_response(ss_raw, "story_summary")
        render_story_summary_docx(ss_data, title, summary_path)

        # Beat Summary
        bs_prompt = build_beat_summary_prompt(synopsis_text)
        bs_raw = call_model(bs_prompt, effective_model, max_tokens=max_tokens)
        bs_data = parse_json_response(bs_raw, "beat_summary")
        render_beat_summary_docx(bs_data, title, beat_path)

        return {
            "status": "success",
            "quality_review_path": quality_path,
            "story_summary_path": summary_path,
            "beat_summary_path": beat_path,
            "error": None,
        }
    except SystemExit:
        return {
            "status": "failure",
            "quality_review_path": quality_path,
            "story_summary_path": summary_path,
            "beat_summary_path": beat_path,
            "error": "LLM call or parse failed — see stderr",
        }


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────


def main() -> int:
    parser = argparse.ArgumentParser(
        description="ANPD V25 synopsis summarizer (v4) — Quality Review + Story Summary."
    )
    parser.add_argument("--synopsis", type=Path, required=True)
    parser.add_argument("--series-config", type=Path, default=None)
    parser.add_argument("--book-config", type=Path, default=None)
    parser.add_argument("--intake", type=Path, default=None)
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--model", type=str, default=None)
    parser.add_argument("--max-tokens", type=int, default=None,
                        help="Override LLM max output tokens per report (default: 8000)")
    args = parser.parse_args()

    if not args.synopsis.exists():
        sys.stderr.write(f"ERROR: Synopsis not found: {args.synopsis}\n")
        return 1
    synopsis_text = args.synopsis.read_text(encoding="utf-8").strip()
    if not synopsis_text:
        sys.stderr.write(f"ERROR: Synopsis is empty: {args.synopsis}\n")
        return 1

    title = resolve_title(args.book_config, args.intake)
    model = resolve_model(args.series_config, args.model)
    output_dir = resolve_output_dir(args.synopsis, args.output_dir)
    effective_max_tokens = args.max_tokens if args.max_tokens else MAX_OUTPUT_TOKENS

    slug = slugify_title(title)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    quality_path = output_dir / f"{slug}_quality_review_{timestamp}.docx"
    summary_path = output_dir / f"{slug}_story_summary_{timestamp}.docx"
    beat_path = output_dir / f"{slug}_beat_summary_{timestamp}.docx"

    sys.stdout.write(f"Synopsis:    {args.synopsis}\n")
    sys.stdout.write(f"Title:       {title}\n")
    sys.stdout.write(f"Model:       {model}\n")
    sys.stdout.write(f"Max tokens:  {effective_max_tokens}\n")
    sys.stdout.write(f"Output 1:    {quality_path}\n")
    sys.stdout.write(f"Output 2:    {summary_path}\n")
    sys.stdout.write(f"Output 3:    {beat_path}\n")

    # ── Report 1: Quality Review ──
    sys.stdout.write("Generating Quality Review...\n")
    qr_prompt = build_quality_review_prompt(synopsis_text)
    qr_raw = call_model(qr_prompt, model, max_tokens=effective_max_tokens)
    qr_data = parse_json_response(qr_raw, "quality_review")
    render_quality_review_docx(qr_data, title, quality_path)

    # ── Report 2: Story Summary ──
    sys.stdout.write("Generating Story Summary...\n")
    ss_prompt = build_story_summary_prompt(synopsis_text)
    ss_raw = call_model(ss_prompt, model, max_tokens=effective_max_tokens)
    ss_data = parse_json_response(ss_raw, "story_summary")
    render_story_summary_docx(ss_data, title, summary_path)

    # ── Report 3: Beat Summary ──
    sys.stdout.write("Generating Beat Summary...\n")
    bs_prompt = build_beat_summary_prompt(synopsis_text)
    bs_raw = call_model(bs_prompt, model, max_tokens=effective_max_tokens)
    bs_data = parse_json_response(bs_raw, "beat_summary")
    render_beat_summary_docx(bs_data, title, beat_path)

    # ── Verify all outputs ──
    for path, label in [
        (quality_path, "Quality Review"),
        (summary_path, "Story Summary"),
        (beat_path, "Beat Summary"),
    ]:
        if not path.exists():
            sys.stderr.write(f"FATAL: {label} docx not written: {path}\n")
            sys.exit(1)
        if path.stat().st_size < 1000:
            sys.stderr.write(
                f"FATAL: {label} docx too small ({path.stat().st_size} bytes): {path}\n"
            )
            sys.exit(1)

    sys.stdout.write(
        f"DONE.\n"
        f"  Quality Review: {quality_path} ({quality_path.stat().st_size} bytes)\n"
        f"  Story Summary:  {summary_path} ({summary_path.stat().st_size} bytes)\n"
        f"  Beat Summary:   {beat_path} ({beat_path.stat().st_size} bytes)\n"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
